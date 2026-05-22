import pdfplumber
import docx
import io


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"PDF read karne mein error: {str(e)}")
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes."""
    text = ""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Tables se bhi text lo
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        raise ValueError(f"DOCX read karne mein error: {str(e)}")
    return text.strip()


def extract_text(uploaded_file) -> str:
    """Main function - file type detect karke text extract karo."""
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif file_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Sirf PDF ya DOCX files supported hain!")
